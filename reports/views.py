from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from django.utils import timezone
from datetime import timedelta
from django.db.models import Sum, Count
from django.http import HttpResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.platypus import PageBreak
from reportlab.lib.enums import TA_CENTER
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from inventory.models import Stock
from sales.models import Sale
from expenses.models import Expense
from django.db.models.functions import TruncDate
from .models import UploadedFile
from .serializers import UploadedFileSerializer, FileUploadSerializer


class ReportViewSet(viewsets.ViewSet):
    """PDF Report generation and file management"""
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    @action(detail=False, methods=['get'])
    def generate(self, request):
        """Generate PDF or DOCX report"""
        report_type = request.query_params.get('period', 'daily')
        format_type = request.query_params.get('format', 'pdf').lower()
        include_charts = request.query_params.get('include_charts', 'true').lower() == 'true'
        include_details = request.query_params.get('include_details', 'true').lower() == 'true'
        start_date_param = request.query_params.get('start_date')
        end_date_param = request.query_params.get('end_date')

        # Determine date range
        if start_date_param and end_date_param:
            try:
                start_date = timezone.datetime.fromisoformat(start_date_param).date()
                end_date = timezone.datetime.fromisoformat(end_date_param).date()
            except ValueError:
                return Response({'error': 'Invalid date format. Use YYYY-MM-DD'}, status=400)
        elif report_type == 'weekly':
            start_date = timezone.now().date() - timedelta(days=7)
            end_date = timezone.now().date()
        elif report_type == 'monthly':
            start_date = timezone.now().date().replace(day=1)
            end_date = timezone.now().date()
        elif report_type == 'yearly':
            start_date = timezone.now().date().replace(month=1, day=1)
            end_date = timezone.now().date()
        else:  # daily
            start_date = timezone.now().date()
            end_date = timezone.now().date()

        # Get data
        sales_query = Sale.objects.filter(
            user=request.user,
            created_at__date__gte=start_date,
            created_at__date__lte=end_date
        )

        expenses_query = Expense.objects.filter(
            user=request.user,
            created_at__date__gte=start_date,
            created_at__date__lte=end_date
        )

        total_income = sales_query.aggregate(Sum('total_amount'))['total_amount__sum'] or 0
        total_expenses = expenses_query.aggregate(Sum('amount'))['amount__sum'] or 0
        total_sales = sales_query.count()
        total_items_sold = sales_query.aggregate(Sum('quantity'))['quantity__sum'] or 0
        total_stocks = Stock.objects.filter(user=request.user).count()

        # Daily breakdown
        daily_sales = sales_query.annotate(date=TruncDate('created_at')).values('date').annotate(
            income=Sum('total_amount'), sales_count=Count('id')
        ).order_by('date')

        daily_expenses = expenses_query.annotate(date=TruncDate('created_at')).values('date').annotate(
            expenses=Sum('amount')
        ).order_by('date')

        # Merge daily data
        daily_data = {}
        for ds in daily_sales:
            daily_data[ds['date']] = {'income': ds['income'], 'sales': ds['sales_count'], 'expenses': 0}
        for de in daily_expenses:
            if de['date'] in daily_data:
                daily_data[de['date']]['expenses'] = de['expenses']
            else:
                daily_data[de['date']] = {'income': 0, 'sales': 0, 'expenses': de['expenses']}

        # Sales details
        sales_details = sales_query.values('stock__name').annotate(
            qty=Sum('quantity'), total=Sum('total_amount')
        ).order_by('-total')

        try:
            if format_type == 'docx':
                return self._generate_docx_report(
                    start_date, end_date, total_income, total_expenses, total_sales,
                    total_items_sold, total_stocks, daily_data, sales_details,
                    include_charts, include_details
                )
            else:
                return self._generate_pdf_report(
                    start_date, end_date, total_income, total_expenses, total_sales,
                    total_items_sold, total_stocks, daily_data, sales_details,
                    include_charts, include_details
                )
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    def _generate_pdf_report(self, start_date, end_date, total_income, total_expenses, total_sales,
                            total_items_sold, total_stocks, daily_data, sales_details,
                            include_charts=True, include_details=True):
        """Generate PDF report"""
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor('#1a365d')
        )

        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=20,
            textColor=colors.HexColor('#2d3748')
        )

        section_style = ParagraphStyle(
            'Section',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=15,
            textColor=colors.HexColor('#2d3748')
        )

        normal_style = styles['Normal']
        normal_style.fontSize = 12
        normal_style.spaceAfter = 12

        # Footer style for page numbers
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray,
            alignment=TA_CENTER
        )

        def on_first_page(canvas, doc):
            # Add footer to first page
            canvas.saveState()
            canvas.setFont('Helvetica', 10)
            canvas.setFillColor(colors.gray)
            canvas.drawString(inch, 0.75 * inch, f"Page 1 of 2 - Generated on {timezone.now().strftime('%-m/%-d/%Y, %-I:%M:%S %p')}")
            canvas.restoreState()

        def on_later_pages(canvas, doc):
            # Add footer to subsequent pages
            canvas.saveState()
            canvas.setFont('Helvetica', 10)
            canvas.setFillColor(colors.gray)
            page_num = canvas.getPageNumber()
            canvas.drawString(inch, 0.75 * inch, f"Page {page_num} of 2 - Generated on {timezone.now().strftime('%-m/%-d/%Y, %-I:%M:%S %p')}")
            canvas.restoreState()

        # Build PDF content
        content = []

        # Title
        content.append(Paragraph("Daily Business Report", title_style))
        content.append(Paragraph(f"Period: {start_date} to {end_date}", subtitle_style))
        content.append(Spacer(1, 20))

        # Summary
        content.append(Paragraph("Summary", section_style))
        summary_data = [
            ['Total Income', f"₹{total_income:,.0f}"],
            ['Total Expenses', f"₹{total_expenses:,.0f}"],
            ['Net Profit', f"₹{(total_income - total_expenses):,.0f}"],
            ['Total Sales', str(total_sales)],
            ['Items Sold', str(total_items_sold)],
            ['Total Stocks', str(total_stocks)],
        ]

        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f7fafc')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1a365d')),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ]))

        content.append(summary_table)
        content.append(Spacer(1, 30))

        # Daily Breakdown Chart (conditional)
        if include_charts and daily_data:
            content.append(Paragraph("Daily Breakdown Chart", section_style))
            drawing = Drawing(400, 200)
            bc = VerticalBarChart()
            bc.x = 50
            bc.y = 50
            bc.height = 125
            bc.width = 300
            bc.data = [
                [d['income'] for d in daily_data.values()],
                [d['expenses'] for d in daily_data.values()]
            ]
            bc.categoryAxis.categoryNames = [d.strftime('%a, %b %d') for d in daily_data.keys()]
            bc.valueAxis.valueMin = 0
            bc.bars[0].fillColor = colors.HexColor('#48bb78')  # green for income
            bc.bars[1].fillColor = colors.HexColor('#f56565')  # red for expenses
            drawing.add(bc)
            content.append(drawing)
            content.append(Spacer(1, 30))

        # Daily Details (conditional)
        if include_details:
            content.append(Paragraph("Daily Details", section_style))
            daily_table_data = [['Date', 'Income', 'Expenses', 'Sales']]
            for date, data in sorted(daily_data.items()):
                daily_table_data.append([
                    date.strftime('%a, %b %d'),
                    f"₹{data['income']:,.0f}",
                    f"₹{data['expenses']:,.0f}",
                    str(data['sales'])
                ])

            daily_table = Table(daily_table_data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1*inch])
            daily_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
                ('FONTSIZE', (0, 1), (-1, -1), 11),
            ]))

            content.append(daily_table)
            content.append(Spacer(1, 30))

        # Sales Details (conditional)
        if include_details:
            content.append(Paragraph("Sales Details", section_style))
            sales_table_data = [['Product', 'Qty', 'Price/Unit', 'Total']]
            for item in sales_details:
                # Get price_per_unit - assuming it's the same for the product
                sale = Sale.objects.filter(
                    stock__name=item['stock__name']
                ).first()
                price_unit = sale.price_per_unit if sale else 0
                sales_table_data.append([
                    item['stock__name'],
                    str(item['qty']),
                    f"₹{price_unit:,.0f}",
                    f"₹{item['total']:,.0f}"
                ])

            sales_table = Table(sales_table_data, colWidths=[2*inch, 1*inch, 1.5*inch, 1.5*inch])
            sales_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
                ('FONTSIZE', (0, 1), (-1, -1), 11),
            ]))

            content.append(sales_table)
            content.append(Spacer(1, 30))

        # Build PDF
        doc.build(content, onFirstPage=on_first_page, onLaterPages=on_later_pages)
        buffer.seek(0)

        # Return PDF response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Daily-Business-Report-{start_date}-to-{end_date}.pdf"'
        return response

    def _generate_docx_report(self, start_date, end_date, total_income, total_expenses, total_sales,
                             total_items_sold, total_stocks, daily_data, sales_details,
                             include_charts=True, include_details=True):
        """Generate DOCX report"""
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Daily Business Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add period
        para = doc.add_paragraph(f'Period: {start_date} to {end_date}')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space
        
        # Summary section
        doc.add_heading('Summary', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Add summary data
        summary_data = [
            ('Total Income', f"₹{total_income:,.0f}"),
            ('Total Expenses', f"₹{total_expenses:,.0f}"),
            ('Net Profit', f"₹{(total_income - total_expenses):,.0f}"),
            ('Total Sales', str(total_sales)),
            ('Items Sold', str(total_items_sold)),
            ('Total Stocks', str(total_stocks)),
        ]
        
        for metric, value in summary_data:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add space
        
        # Daily Details section (conditional)
        if include_details:
            doc.add_heading('Daily Details', level=1)
            daily_table = doc.add_table(rows=1, cols=4)
            daily_table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = daily_table.rows[0].cells
            hdr_cells[0].text = 'Date'
            hdr_cells[1].text = 'Income'
            hdr_cells[2].text = 'Expenses'
            hdr_cells[3].text = 'Sales'
            
            # Add daily data
            for date, data in sorted(daily_data.items()):
                row_cells = daily_table.add_row().cells
                row_cells[0].text = date.strftime('%a, %b %d')
                row_cells[1].text = f"₹{data['income']:,.0f}"
                row_cells[2].text = f"₹{data['expenses']:,.0f}"
                row_cells[3].text = str(data['sales'])
            
            doc.add_paragraph()  # Add space
        
        # Sales Details section (conditional)
        if include_details:
            doc.add_heading('Sales Details', level=1)
            sales_table = doc.add_table(rows=1, cols=4)
            sales_table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = sales_table.rows[0].cells
            hdr_cells[0].text = 'Product'
            hdr_cells[1].text = 'Qty'
            hdr_cells[2].text = 'Price/Unit'
            hdr_cells[3].text = 'Total'
            
            # Add sales data
            for item in sales_details:
                sale = Sale.objects.filter(stock__name=item['stock__name']).first()
                price_unit = sale.price_per_unit if sale else 0
                
                row_cells = sales_table.add_row().cells
                row_cells[0].text = item['stock__name']
                row_cells[1].text = str(item['qty'])
                row_cells[2].text = f"₹{price_unit:,.0f}"
                row_cells[3].text = f"₹{item['total']:,.0f}"
            
            doc.add_paragraph()  # Add space
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated on {timezone.now().strftime('%-m/%-d/%Y, %-I:%M:%S %p')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return DOCX response
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Daily-Business-Report-{start_date}-to-{end_date}.docx"'
        return response

    @action(detail=False, methods=['post'])
    def upload_file(self, request):
        """Upload PDF or DOCX file"""
        try:
            serializer = FileUploadSerializer(data=request.data)
            if serializer.is_valid():
                file = serializer.validated_data['file']
                
                # Create UploadedFile instance
                uploaded_file = UploadedFile.objects.create(
                    user=request.user,
                    file=file,
                    filename=file.name,
                    file_type=file.name.lower().split('.')[-1],
                    file_size=file.size
                )
                
                # Return serialized data
                response_serializer = UploadedFileSerializer(uploaded_file)
                return Response(response_serializer.data, status=201)
            else:
                return Response(serializer.errors, status=400)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    @action(detail=False, methods=['get'])
    def list_files(self, request):
        """List uploaded files for the current user"""
        try:
            files = UploadedFile.objects.filter(user=request.user)
            serializer = UploadedFileSerializer(files, many=True)
            return Response(serializer.data)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    @action(detail=False, methods=['delete'])
    def delete_file(self, request):
        """Delete an uploaded file"""
        file_id = request.query_params.get('file_id')
        if not file_id:
            return Response({'error': 'file_id parameter is required'}, status=400)
        
        try:
            file = UploadedFile.objects.get(id=file_id, user=request.user)
            # Delete the file from storage
            if file.file:
                file.file.delete()
            file.delete()
            return Response({'message': 'File deleted successfully'})
        except UploadedFile.DoesNotExist:
            return Response({'error': 'File not found'}, status=404)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    @action(detail=False, methods=['get'])
    def download_file(self, request):
        """Download an uploaded file"""
        file_id = request.query_params.get('file_id')
        if not file_id:
            return Response({'error': 'file_id parameter is required'}, status=400)
        
        try:
            file = UploadedFile.objects.get(id=file_id, user=request.user)
            
            # Determine content type
            if file.file_type.lower() == 'pdf':
                content_type = 'application/pdf'
            else:
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Return file response
            response = HttpResponse(file.file.read(), content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{file.filename}"'
            return response
        except UploadedFile.DoesNotExist:
            return Response({'error': 'File not found'}, status=404)
        except Exception as e:
            return Response({'error': str(e)}, status=500)
